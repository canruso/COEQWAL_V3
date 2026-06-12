from __future__ import annotations

import os
from collections import defaultdict

# python-docx is only needed by the legacy generate_scenario_review_doc()
# function below. Imported lazily to avoid forcing it as a dependency.

from coeqwalpackage.review_config import (
    build_doc_spec_with_labels,
    get_subdir,
    build_filename,
    parse_scenario_set_configs,
)

_SECTION_PATTERNS = [
    ("Reservoir Storage",  lambda v: v.startswith("S_")),
    ("Deliveries",         lambda v: v.startswith("DEL_") or v.startswith("C_DMC") or v.startswith("C_CAA")),
    ("Flows & Salinity",   lambda v: v.startswith("C_SAC") or v.startswith("C_SJR") or v.startswith("SP_SAC")
                                     or v.startswith("X2_") or v.endswith("_EC_MONTH")),
    ("Stream Gain",        lambda v: v.startswith("SG_")),
    ("Applied Water",      lambda v: v.startswith("AWO")),
]

_SECTION_ORDER = [
    "Reservoir Storage",
    "Deliveries",
    "Flows & Salinity",
    "Stream Gain",
    "Applied Water",
    "Other",
]


def _get_section(varname: str) -> str:
    for section_name, test in _SECTION_PATTERNS:
        if test(varname):
            return section_name
    return "Other"


def _add_picture_safe(doc, path: str, width_inches: float, required: bool = True) -> bool:
    from docx.shared import Inches
    if os.path.isfile(path):
        doc.add_picture(path, width=Inches(width_inches))
        return True
    if required:
        print(f"  [MISSING] {path}")
    return False


def list_available_scenario_sets(plots_root: str, scenario_groupings_csv: str):
    return parse_scenario_set_configs(plots_root, scenario_groupings_csv)


def generate_scenario_review_doc(
    plots_base: str,
    set_name: str,
    baseline: int,
    compare: list[int],
    output_path: str,
    *,
    width_inches: float = 7.0,
    placeholder: str = "{{Add analysis here}}",
    summary_placeholder: str = (
        "{{After adding and reviewing plots, summarize the outcomes in this "
        "scenario compared to the baseline. Note major differences, patterns, "
        "and anything unexpected.}}"
    ),
):
    from docx import Document
    from docx.shared import Inches

    if not os.path.isdir(plots_base):
        raise FileNotFoundError(f"Plots folder not found: {plots_base}")

    doc_spec = build_doc_spec_with_labels(plots_base)
    if not doc_spec or not doc_spec[0]["variables"]:
        print(f"[SKIP] No plot files found under {plots_base}")
        return None

    scenario_id_str = ", ".join(f"s{s:04d}" for s in compare)
    baseline_str = f"s{baseline:04d}"

    doc = Document()

    doc.add_heading(
        f"CalSim3 Scenario Output Review: {scenario_id_str} vs {baseline_str}",
        level=0,
    )
    doc.add_paragraph(f"Scenario ID(s): {scenario_id_str}")
    doc.add_paragraph(f"Baseline: {baseline_str}")
    doc.add_paragraph("Scenario Description(s):")
    doc.add_paragraph("Reference Description:")
    doc.add_paragraph("Review Date (updates appended):")
    doc.add_paragraph("Reviewer(s):")
    doc.add_paragraph(f"Summary of Findings:\n{summary_placeholder}")

    variables = doc_spec[0]["variables"]

    sections_dict = defaultdict(list)
    for var_entry in variables:
        sections_dict[_get_section(var_entry["varname"])].append(var_entry)

    for sec_idx, section_name in enumerate(_SECTION_ORDER, start=1):
        if section_name not in sections_dict:
            continue

        doc.add_heading(f"{sec_idx}. {section_name}", level=1)

        for var_idx, var_entry in enumerate(sections_dict[section_name], start=1):
            varname = var_entry["varname"]
            label = var_entry["label"]
            plot_types = var_entry["plots"]

            doc.add_heading(f"{var_idx}. {label}", level=2)

            any_added = False
            for plot_type in plot_types:
                subdir = get_subdir(plot_type)
                filename = build_filename(varname, plot_type)
                img_path = os.path.join(plots_base, subdir, filename)
                is_tucp = "tucp" in plot_type.lower()

                if is_tucp:
                    if os.path.isfile(img_path):
                        doc.add_picture(img_path, width=Inches(width_inches))
                        doc.add_paragraph(placeholder)
                        any_added = True
                else:
                    if _add_picture_safe(doc, img_path, width_inches, required=True):
                        doc.add_paragraph(placeholder)
                        any_added = True

            if not any_added:
                doc.add_paragraph(f"[No plot files found for {varname}]")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


def run_scenario_review(
    plots_root: str,
    review_output_dir: str,
    scenario_groupings_csv: str,
    *,
    width_inches: float = 7.0,
    placeholder: str = "{{Add analysis here}}",
    summary_placeholder: str = (
        "{{After adding and reviewing plots, summarize the outcomes in this "
        "scenario compared to the baseline. Note major differences, patterns, "
        "and anything unexpected.}}"
    ),
    selected_set_name: str | None = None,
    file_prefix: str = "Scenario_Review",
) -> list[str]:
    scenario_set_configs = parse_scenario_set_configs(plots_root, scenario_groupings_csv)

    if selected_set_name is not None:
        scenario_set_configs = [
            cfg for cfg in scenario_set_configs
            if cfg["set_name"] == selected_set_name
        ]

    if not scenario_set_configs:
        raise FileNotFoundError(
            "No valid scenario sets found. Check scenario_groupings.csv and plots_root."
        )

    os.makedirs(review_output_dir, exist_ok=True)
    saved_paths = []

    for cfg in scenario_set_configs:
        set_name = cfg["set_name"]
        plots_base = cfg["plots_base"]

        print(f"\nGenerating review template for: {set_name}")

        out_filename = f"{file_prefix}_{set_name}_Not_Annotated.docx"
        out_path = os.path.join(review_output_dir, out_filename)

        saved = generate_scenario_review_doc(
            plots_base=plots_base,
            set_name=set_name,
            baseline=cfg["baseline"],
            compare=cfg["compare"],
            output_path=out_path,
            width_inches=width_inches,
            placeholder=placeholder,
            summary_placeholder=summary_placeholder,
        )

        if saved is not None:
            saved_paths.append(saved)

    return saved_paths

def run_batch_review(
    df,
    plots_root: str,
    set_name: str,
    all_scenarios: list,
    baseline: int,
    scenario_labels: dict,
    output_base: str,
    *,
    doc_spec=None,
    scenario_context=None,
    set_context=None,
    tucp_years=None,
    wyt_wet=None,
    wyt_dry=None,
    wyt_month=5,
    sections=None,
    variables=None,
    plot_types=None,
    dry_run=False,
) -> dict:
    """Batch-process all plots for one scenario set through the LLM pipeline.

    For each (variable, plot_type) entry in doc_spec, this function:
      1. Checks whether a JSON output already exists (skips if so).
      2. Checks whether the image file exists on disk (skips if not).
      3. Computes statistics from df via compute_var_stats().
      4. Builds a prompt via build_prompt().
      5. Calls analyze_plot() to get the LLM observation.
      6. Saves the result as a JSON file under output_base/set_name/<subdir>/.

    Parameters
    ----------
    df : pd.DataFrame
        The full multi-scenario dataframe loaded by read_in_df().
    plots_root : str
        Root directory containing one subfolder per scenario set.
    set_name : str
        Name of the scenario set folder (e.g. "s20_s35_s36_s37").
    all_scenarios : list of int
        All scenario IDs for this set, baseline first.
    baseline : int
        Baseline scenario ID.
    scenario_labels : dict
        Optional mapping of scenario ID -> human label. Pass {} if unused.
    output_base : str
        Root directory where JSON outputs are written.
    doc_spec : list or None
        Pre-built doc spec from build_doc_spec_with_labels(). If None, it is
        built dynamically from the plots folder.
    scenario_context : dict or None
        Legacy simple {sid: description} mapping passed to build_prompt().
    set_context : dict or None
        Rich scenario set context from load_set_context(). Takes precedence
        over scenario_context. Has keys: description, scenarios, expectations,
        questions. Injected into every per-plot prompt.
    tucp_years : dict or None
        Mapping of scenario ID -> list of TUCP water years.
    wyt_wet, wyt_dry : list or None
        Water year type lists for wet/dry filtering.
    wyt_month : int
        Month used to assign water year types (default 5 = May).
    sections, variables, plot_types : list or None
        Optional filters to restrict processing to a subset.
    dry_run : bool
        If True, enumerate what would be processed without calling the LLM.

    Returns
    -------
    dict with keys:
        "processed" : list of dicts describing items that were (or would be) processed.
        "skipped"   : list of dicts describing items that were skipped and why.
    """
    import json as _json
    from coeqwalpackage.review_config import (
        build_doc_spec_with_labels, iter_doc_spec, get_subdir, build_filename, get_units,
    )
    from coeqwalpackage.llm_utils import analyze_plot, compute_var_stats
    from coeqwalpackage.prompts import build_prompt

    set_plots_dir = os.path.join(plots_root, set_name)

    if doc_spec is None:
        doc_spec = build_doc_spec_with_labels(set_plots_dir)

    processed = []
    skipped = []

    for section_title, varname, label, units, plot_type in iter_doc_spec(doc_spec):
        # Apply optional filters
        if sections is not None and section_title not in sections:
            skipped.append({"tag": f"{varname}/{plot_type}", "reason": "section filtered"})
            continue
        if variables is not None and not any(varname.startswith(v) for v in variables):
            skipped.append({"tag": f"{varname}/{plot_type}", "reason": "variable filtered"})
            continue
        if plot_types is not None and plot_type not in plot_types:
            skipped.append({"tag": f"{varname}/{plot_type}", "reason": "plot_type filtered"})
            continue

        subdir = get_subdir(plot_type)
        img_filename = build_filename(varname, plot_type)
        img_path = os.path.join(set_plots_dir, subdir, img_filename)
        json_path = os.path.join(
            output_base, set_name, subdir,
            img_filename.replace(".png", ".json"),
        )

        # Skip if image does not exist (e.g. TUCP plots with no TUCP data)
        if not os.path.isfile(img_path):
            skipped.append({"tag": f"{varname}/{plot_type}", "reason": "image missing"})
            continue

        # Skip if JSON already exists (resume support)
        if os.path.isfile(json_path):
            skipped.append({"tag": f"{varname}/{plot_type}", "reason": "already done"})
            continue

        entry = {
            "tag": f"{varname}/{plot_type}",
            "varname": varname,
            "var_label": label,
            "plot_type": plot_type,
            "img_path": img_path,
            "json_path": json_path,
        }
        processed.append(entry)

        if dry_run:
            continue

        # Compute stats
        stats_text = compute_var_stats(
            df=df,
            varname=varname,
            units=units,
            scenarios=all_scenarios,
            baseline=baseline,
            plot_type=plot_type,
            scenario_labels=scenario_labels or {},
            tucp_years=tucp_years,
            wyt_wet=wyt_wet,
            wyt_dry=wyt_dry,
            wyt_month=wyt_month,
        )

        # Build scenario name strings for prompt
        scenario_names = ", ".join(
            scenario_labels.get(s, f"s{s:04d}") for s in all_scenarios if s != baseline
        )
        baseline_name = scenario_labels.get(baseline, f"s{baseline:04d}")

        prompt = build_prompt(
            plot_type=plot_type,
            var_label=label,
            scenario_names=scenario_names,
            baseline=baseline_name,
            scenario_context=scenario_context,
            stats_text=stats_text if stats_text else None,
            set_context=set_context,
        )

        # Call LLM
        try:
            result = analyze_plot(
                image_path=img_path,
                prompt=prompt,
                stats_text=stats_text,
                scenarios=all_scenarios,
                baseline=baseline,
            )
        except Exception as exc:
            print(f"  [ERROR] {varname}/{plot_type}: {exc}")
            skipped.append({"tag": f"{varname}/{plot_type}", "reason": f"LLM error: {exc}"})
            processed.pop()
            continue

        # Save JSON output
        os.makedirs(os.path.dirname(json_path), exist_ok=True)
        output_data = {
            "var_label":   label,
            "plot_type":   plot_type,
            "stats":       stats_text,
            "observation": result.get("narrative", ""),
            "structured":  result.get("structured", {}),
            "validation":  result.get("validation", []),
            "raw":         result.get("raw", ""),
        }
        with open(json_path, "w") as fh:
            _json.dump(output_data, fh, indent=2)

        print(f"  [OK] {varname}/{plot_type} -> {os.path.relpath(json_path, output_base)}")

    return {"processed": processed, "skipped": skipped}
