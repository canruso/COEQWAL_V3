import os
from docx import Document # pip install python-docx
from docx.shared import Inches
from coeqwalpackage import scenario_review
document = Document()

def run_scenario_review(all_scenarios):
    for scenarios in all_scenarios:
        # Pad with leading zeros (s0020, s0039, ...)
        scenario_tags = [f"s{n}" for n in scenarios]
        
        # Baseline scenario ID: all scenarios joined with underscores
        baseline_scenario_id = "_".join(scenario_tags)
        
        # Scenario ID: everything except the baseline
        scenario_id = ", ".join(scenario_tags[1:])
    
        file_path = f'../../CalSim3_Model_Runs/Scenarios/Group_Data_Extraction/plots_output/{baseline_scenario_id}'
    
        if os.path.exists(file_path):
            print(f'{baseline_scenario_id} found')
        else:
            print(f'{baseline_scenario_id} not found')
            continue
    
        mon_ts_path = f'../../CalSim3_Model_Runs/Scenarios/Group_Data_Extraction/plots_output/{baseline_scenario_id}/mon_ts/'
        moy_avg_path = f'../../CalSim3_Model_Runs/Scenarios/Group_Data_Extraction/plots_output/{baseline_scenario_id}/moy_avg/'
        ann_exceed_path = f'../../CalSim3_Model_Runs/Scenarios/Group_Data_Extraction/plots_output/{baseline_scenario_id}/ann_exceed/'
        ann_tot_path = f'../../CalSim3_Model_Runs/Scenarios/Group_Data_Extraction/plots_output/{baseline_scenario_id}/ann_tot_path/'
        exceedance_path = f'../../CalSim3_Model_Runs/Scenarios/Group_Data_Extraction/plots_output/{baseline_scenario_id}/exceedance_path/'
        
        # Add TUCP plots only if they exist
        def tucp_check(document, file_path, description="{{example description/interpretation}}", width=Inches(7.0)):
            if os.path.exists(file_path):
                document.add_picture(file_path, width=width)
                document.add_paragraph(description)
        
        document.add_heading(f'CalSim3 Scenario Output Review: {scenario_id}', level=0)
        
        document.add_paragraph(f"Scenario ID(s): {scenario_id}")
        document.add_paragraph("Scenario Description(s):")
        document.add_paragraph("Reference Description:")
        document.add_paragraph("Review Date (updates appended):")
        document.add_paragraph("Reviewers(s):")
        document.add_paragraph("Summary of Findings:\n{{after adding and reviewing plots, how would you summarize the outcomes in this scenario compared ot the baseline? Are certain outcome types consistently higher or lower? Do some outcomes not make sense?}}")
        
        document.add_heading('1. Reservoir Storage', level=1)
        
        document.add_heading('1. Shasta Reservoir', level=2)
        
        document.add_picture(f'{mon_ts_path}S_SHSTA__mon_ts.png', width=Inches(7.0)) # Monthly timeseries
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}S_SHSTA__ann_exceed_apr.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_SHSTA__ann_exceed_apr_tucp.png') # April exceedance (TUCP)
        document.add_picture(f'{ann_exceed_path}S_SHSTA__ann_exceed_sept.png', width=Inches(7.0)) # September exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_SHSTA__ann_exceed_sept_tucp.png') # September exceedance (TUCP)
        document.add_picture(f'{moy_avg_path}S_SHSTA__moy_all.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}S_SHSTA__moy_tucp.png') # MOY average (TUCP)
        document.add_paragraph("{{example description/interpretation}}")
        
        document.add_heading('2. Oroville Reservoir', level=2)
        
        document.add_picture(f'{mon_ts_path}S_OROVL__mon_ts.png', width=Inches(7.0)) # Monthly timeseries
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}S_OROVL__ann_exceed_apr.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_OROVL__ann_exceed_apr_tucp.png') # April exceedance (TUCP)
        document.add_picture(f'{ann_exceed_path}S_OROVL__ann_exceed_sept.png', width=Inches(7.0)) # September exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_OROVL__ann_exceed_sept_tucp.png') # September exceedance (TUCP)
        document.add_picture(f'{moy_avg_path}S_OROVL__moy_all.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}S_OROVL__moy_tucp.png') # MOY average (TUCP)
        document.add_paragraph("{{example description/interpretation}}")
        
        document.add_heading('3. Folsom Reservoir', level=2)
        
        document.add_picture(f'{mon_ts_path}S_FOLSM__mon_ts.png', width=Inches(7.0)) # Monthly timeseries
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}S_FOLSM__ann_exceed_apr.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_FOLSM__ann_exceed_apr_tucp.png') # April exceedance (TUCP)
        document.add_picture(f'{ann_exceed_path}S_FOLSM__ann_exceed_sept.png', width=Inches(7.0)) # September exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_FOLSM__ann_exceed_sept_tucp.png') # September exceedance (TUCP)
        document.add_picture(f'{moy_avg_path}S_FOLSM__moy_all.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}S_FOLSM__moy_tucp.png') # MOY average (TUCP)
        document.add_paragraph("{{example description/interpretation}}")
        
        document.add_heading('4. New Melones Reservoir', level=2)
        
        document.add_picture(f'{mon_ts_path}S_MELON__mon_ts.png', width=Inches(7.0)) # Monthly timeseries
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}S_MELON__ann_exceed_apr.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_MELON__ann_exceed_apr_tucp.png') # April exceedance (TUCP)
        document.add_picture(f'{ann_exceed_path}S_MELON__ann_exceed_sept.png', width=Inches(7.0)) # September exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_MELON__ann_exceed_sept_tucp.png') # September exceedance (TUCP)
        document.add_picture(f'{moy_avg_path}S_MELON__moy_all.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}S_MELON__moy_tucp.png') # MOY average (TUCP)
        document.add_paragraph("{{example description/interpretation}}")
        
        document.add_heading('5. San Luis Reservoir', level=2)
        
        document.add_picture(f'{mon_ts_path}S_SLUIS_s_mon_ts.png', width=Inches(7.0)) # Monthly timeseries
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}S_SLUIS_s_ann_exceed_apr.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_SLUIS_s_ann_exceed_apr_tucp.png') # April exceedance (TUCP)
        document.add_picture(f'{ann_exceed_path}S_SLUIS_s_ann_exceed_sept.png', width=Inches(7.0)) # September exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}S_SLUIS_s_ann_exceed_sept_tucp.png') # September exceedance (TUCP)
        document.add_picture(f'{moy_avg_path}S_SLUIS_s_moy_all.png', width=Inches(7.0)) # April exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}S_SLUIS_s_moy_tucp.png') # MOY average (TUCP)
        document.add_paragraph("{{example description/interpretation}}")
        
        document.add_heading('2. Deliveries', level=1)
        
        document.add_heading('1. CVP Project Agriculture - NOD', level=2)
        
        document.add_picture(f'{moy_avg_path}DEL_CVP_PAG_N_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}DEL_CVP_PAG_N_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}DEL_CVP_PAG_N_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}DEL_CVP_PAG_N_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('2. CVP Project Agriculture - SOD', level=2)
        
        document.add_picture(f'{moy_avg_path}DEL_CVP_PAG_S_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}DEL_CVP_PAG_S_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}DEL_CVP_PAG_S_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}DEL_CVP_PAG_S_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('3. CVP Settlement Contractors - NOD', level=2)
        
        document.add_picture(f'{moy_avg_path}DEL_CVP_PSC_N_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}DEL_CVP_PSC_N_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}DEL_CVP_PSC_N_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}DEL_CVP_PSC_N_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('4. CVP Exchance Contractors - SOD', level=2)
        
        document.add_picture(f'{moy_avg_path}DEL_CVP_PEX_S_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}DEL_CVP_PEX_S_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}DEL_CVP_PEX_S_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}DEL_CVP_PEX_S_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('5. SWP Total Table A Contractors', level=2)
        
        document.add_picture(f'{moy_avg_path}DEL_SWP_TOTA__moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}DEL_CVP_SWP_TOTA__moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}DEL_SWP_TOTA__ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}DEL_SWP_TOTA__ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('6. SWP Municipal - SOD', level=2)
        
        document.add_picture(f'{moy_avg_path}DEL_SWP_PMI_S_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}DEL_CVP_SWP_PMI_S_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}DEL_SWP_PMI_S_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}DEL_SWP_PMI_S_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('7. CVP South Delta Exports', level=2)
        
        document.add_picture(f'{moy_avg_path}C_DMC000_TD_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_DMC000_TD_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}C_DMC000_TD_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_DMC000_TD_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('8. SWP South Delta Exports', level=2)
        
        document.add_picture(f'{moy_avg_path}C_CAA003_TD_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_CAA003_TD_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{ann_exceed_path}C_CAA003_TD_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_CAA003_TD_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('3. Flows & Salinity', level=1)
        
        document.add_heading('1. Delta Outflow', level=2)
        
        document.add_picture(f'{moy_avg_path}C_SAC000_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_SAC000_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}C_SAC000_s_moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}C_SAC000_s_moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}C_SAC000_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_SAC000_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('2. Sacramento River Inflow to the Delta', level=2)
        
        document.add_picture(f'{moy_avg_path}C_SAC041_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_SAC041_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}C_SAC041_s_moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}C_SAC041_s_moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}C_SAC041_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_SAC041_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_picture(f'{moy_avg_path}C_SAC085_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_SAC085_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}C_SAC085_s_moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}C_SAC085_s_moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}C_SAC085_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_SAC085_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_picture(f'{moy_avg_path}C_SAC122_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_SAC122_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}C_SAC122_s_moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}C_SAC122_s_moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}C_SAC122_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_SAC122_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_picture(f'{moy_avg_path}SP_SAC083_YBP037_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}SP_SAC083_YBP037_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}SP_SAC083_YBP037_moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}SP_SAC083_YBP037_moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}SP_SAC083_YBP037_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}SP_SAC083_YBP037_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('3. San Joaquin River Inflow to the Delta', level=2)
        
        document.add_picture(f'{moy_avg_path}C_SJR070_s_moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}C_SJR070_s_moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}C_SJR070_s_moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}C_SJR070_s_moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}C_SJR070_s_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}C_SJR070_s_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('4. Salinity', level=2)
        
        document.add_picture(f'{moy_avg_path}X2_PRV_KM__moy_all.png', width=Inches(7.0)) # MOY average
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{moy_avg_path}X2_PRV_KM__moy_tucp.png') # MOY average (TUCP)
        document.add_picture(f'{moy_avg_path}X2_PRV_KM__moy_dry.png', width=Inches(7.0)) # MOY average (dry)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{moy_avg_path}X2_PRV_KM__moy_wet.png', width=Inches(7.0)) # MOY average (wet)
        document.add_paragraph("{{example description/interpretation}}")
        document.add_picture(f'{ann_exceed_path}X2_PRV_KM__ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
        document.add_paragraph("{{example description/interpretation}}")
        tucp_check(document, f'{ann_exceed_path}X2_PRV_KM__ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('4. Stream Gain', level=1)
        
        document.add_heading('1. SG_SACAB', level=2)
        
        if not os.path.exists(f'{moy_avg_path}SG_SACAB_moy_all.png'):
            print("Missing variable SG_SACAB")
        else:
            document.add_picture(f'{moy_avg_path}SG_SACAB_moy_all.png', width=Inches(7.0)) # MOY average
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{moy_avg_path}SG_SACAB_moy_tucp.png') # MOY average (TUCP)
            document.add_picture(f'{ann_exceed_path}SG_SACAB_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{ann_exceed_path}SG_SACAB_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('2. SG_SACBB', level=2)
        
        if not os.path.exists(f'{moy_avg_path}SG_SACBB_moy_all.png'):
            print("Missing variable SG_SACBB")
        else:
            document.add_picture(f'{moy_avg_path}SG_SACBB_moy_all.png', width=Inches(7.0)) # MOY average
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{moy_avg_path}SG_SACBB_moy_tucp.png') # MOY average (TUCP)
            document.add_picture(f'{ann_exceed_path}SG_SACBB_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{ann_exceed_path}SG_SACBB_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('3. SG_SACFB', level=2)
        
        if not os.path.exists(f'{moy_avg_path}SG_SACFB_moy_all.png'):
            print("Missing variable SG_SACFB")
        else:
            document.add_picture(f'{moy_avg_path}SG_SACFB_moy_all.png', width=Inches(7.0)) # MOY average
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{moy_avg_path}SG_SACFB_moy_tucp.png') # MOY average (TUCP)
            document.add_picture(f'{ann_exceed_path}SG_SACFB_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{ann_exceed_path}SG_SACFB_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('4. SG_SACAMR', level=2)
        
        if not os.path.exists(f'{moy_avg_path}SG_SACAMR_moy_all.png'):
            print("Missing variable SG_SACAMR")
        else:
            document.add_picture(f'{moy_avg_path}SG_SACAMR_moy_all.png', width=Inches(7.0)) # MOY average
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{moy_avg_path}SG_SACAMR_moy_tucp.png') # MOY average (TUCP)
            document.add_picture(f'{ann_exceed_path}SG_SACAMR_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{ann_exceed_path}SG_SACAMR_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        document.add_heading('5. SG_SACBASIN', level=2)
        
        if not os.path.exists(f'{moy_avg_path}SG_SACBASIN_moy_all.png'):
            print("Missing variable SG_SACBASIN")
        else:
            document.add_picture(f'{moy_avg_path}SG_SACBASIN_moy_all.png', width=Inches(7.0)) # MOY average
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{moy_avg_path}SG_SACBASIN_moy_tucp.png') # MOY average (TUCP)
            document.add_picture(f'{ann_exceed_path}SG_SACBASIN_ann_exceed_all.png', width=Inches(7.0)) # Annual exceedance
            document.add_paragraph("{{example description/interpretation}}")
            tucp_check(document, f'{ann_exceed_path}SG_SACBASIN_ann_exceed_tucp.png') # Annual exceedance (TUCP)
        
        review_dir = "../../CalSim3_Model_Runs/Scenarios/Performance_Metrics/Scenario_Review"
        if not os.path.exists(review_dir):
            os.makedirs(review_dir)
        
        review_path = os.path.join(review_dir, f'Scenarios_{scenarios}_Review_Not_Annotated.docx')
        document.save(review_path)
        print("Saved scenario review template in:")
        print(review_path)